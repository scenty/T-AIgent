# -*- coding: utf-8 -*-
"""创建 Word 简报模板"""

from pathlib import Path
from docx import Document

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "report" / "templates"


def main():
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("{{ title }}", 0)
    doc.add_paragraph("区域: {{ region }}")
    doc.add_paragraph("预报时效: {{ time_label }}")
    doc.add_paragraph("生成时间: {{ generated_at }}")
    doc.add_heading("综合风险等级", level=1)
    doc.add_paragraph("{{ comprehensive_level }}")
    doc.add_heading("海况概况", level=1)
    doc.add_paragraph("{% for a in assessments %}{{ a.element }}: {{ a.value }} {{ a.unit }} -> {{ a.level_name }}\n{% endfor %}")
    doc.add_heading("影响分析", level=1)
    doc.add_paragraph("{{ impact_analysis }}")
    doc.add_heading("建议措施", level=1)
    doc.add_paragraph("{{ advice }}")
    doc.add_paragraph("{{ disclaimer }}")
    doc.save(str(TEMPLATE_DIR / "briefing.docx"))
    print(f"模板已创建: {TEMPLATE_DIR / 'briefing.docx'}")


if __name__ == "__main__":
    main()
