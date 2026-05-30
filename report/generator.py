# -*- coding: utf-8 -*-
"""海洋灾害风险简报生成"""

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

from jinja2 import Environment, FileSystemLoader
from docx import Document

from risk.assessor import session_context

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_DIR = BASE_DIR / "report" / "templates"
OUTPUT_DIR = BASE_DIR / "output"


def _ensure_dirs():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _build_context(
    title: Optional[str] = None,
    region: Optional[str] = None,
    data_sources: Optional[list] = None,
    stats_summary: Optional[list] = None,
    risk_data: Optional[dict] = None,
    impact_analysis: Optional[str] = None,
) -> dict:
    risk = risk_data or session_context.get("last_risk") or {}
    stats = stats_summary or session_context.get("last_stats") or []

    comp = risk.get("comprehensive", {})
    assessments = risk.get("assessments", [])

    return {
        "title": title or f"{region or '目标海域'}海洋灾害风险简报",
        "region": region or risk.get("location", "未指定"),
        "time_label": risk.get("time", datetime.now().strftime("%Y-%m-%d %H:%M")),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "data_sources": data_sources or [],
        "stats_summary": stats,
        "assessments": assessments,
        "comprehensive_level": comp.get("level_name", "未评估"),
        "comprehensive_level_id": comp.get("level_id", "none"),
        "advice": comp.get("advice", ""),
        "impact_analysis": impact_analysis or _default_impact(comp.get("level_id", "none"), region),
        "disclaimer": "本简报由海洋预报 Agent 自动生成，仅供参考，请以官方发布为准。",
    }


def _default_impact(level_id: str, region: Optional[str]) -> str:
    region = region or "目标海域"
    impacts = {
        "none": f"{region}海况总体平稳，适宜正常海上活动。",
        "blue": f"{region}海况有所增强，小型船舶需注意安全。",
        "yellow": f"{region}海况较差，建议停止近海养殖和施工作业。",
        "orange": f"{region}海况恶劣，应停止一切海上作业并组织船舶回港。",
        "red": f"{region}海况极其危险，应立即启动最高级别应急响应。",
    }
    return impacts.get(level_id, f"{region}海况需关注。")


def generate_briefing_preview(
    title: Optional[str] = None,
    region: Optional[str] = None,
    data_sources: Optional[list] = None,
    stats_summary: Optional[list] = None,
    risk_data: Optional[dict] = None,
    impact_analysis: Optional[str] = None,
) -> Dict[str, Any]:
    _ensure_dirs()
    ctx = _build_context(title, region, data_sources, stats_summary, risk_data, impact_analysis)

    env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)), autoescape=False)
    template = env.get_template("briefing.md.j2")
    content = template.render(**ctx)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = OUTPUT_DIR / f"briefing_{ts}.md"
    out_path.write_text(content, encoding="utf-8")

    return {
        "success": True,
        "format": "markdown",
        "path": str(out_path),
        "preview": content,
    }


def generate_briefing_docx(
    title: Optional[str] = None,
    region: Optional[str] = None,
    data_sources: Optional[list] = None,
    stats_summary: Optional[list] = None,
    risk_data: Optional[dict] = None,
    impact_analysis: Optional[str] = None,
) -> Dict[str, Any]:
    _ensure_dirs()
    ctx = _build_context(title, region, data_sources, stats_summary, risk_data, impact_analysis)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = OUTPUT_DIR / f"briefing_{ts}.docx"

    template_path = TEMPLATE_DIR / "briefing.docx"
    if template_path.exists():
        from docxtpl import DocxTemplate
        doc = DocxTemplate(str(template_path))
        doc.render(ctx)
        doc.save(str(out_path))
    else:
        doc = Document()
        doc.add_heading(ctx["title"], 0)
        doc.add_paragraph(f"区域: {ctx['region']}")
        doc.add_paragraph(f"时间: {ctx['time_label']}")
        doc.add_paragraph(f"综合风险等级: {ctx['comprehensive_level']}")
        doc.add_heading("海况概况", level=1)
        for a in ctx["assessments"]:
            doc.add_paragraph(f"{a.get('element')}: {a.get('value')} {a.get('unit')} -> {a.get('level_name')}")
        doc.add_heading("影响分析", level=1)
        doc.add_paragraph(ctx["impact_analysis"])
        doc.add_heading("建议措施", level=1)
        doc.add_paragraph(ctx["advice"])
        doc.add_paragraph(ctx["disclaimer"])
        doc.save(str(out_path))

    return {
        "success": True,
        "format": "docx",
        "path": str(out_path),
    }


def list_briefing_templates() -> Dict[str, Any]:
    templates = []
    if TEMPLATE_DIR.exists():
        for f in TEMPLATE_DIR.iterdir():
            templates.append({"name": f.name, "path": str(f)})
    return {"success": True, "templates": templates}
